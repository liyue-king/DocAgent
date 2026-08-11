"""
====================================================================
文件用途：行业知识库 RAG 服务（docagent_knowledge 向量集合）
====================================================================
作用：
    1. 行业→模板指南的向量化存储与语义检索（供聊天机器人回答
       “哪个行业用哪个模板”）。
    2. 支持上传自定义文档（txt/md/docx）：文本抽取 → 段落切块 →
       BGE-M3 向量化 → 灌入 ChromaDB；原文同时备份到 MinIO。
依赖：
    - chromadb（向量库 HTTP 客户端）
    - app.services.embeddings.get_embedder（共享 BGE-M3 单例）
    - app.services.storage.storage（MinIO 原文备份）
调用方：
    - app/api/chat.py（聊天检索 / 上传 / 统计）
    - scripts/init_knowledge.py（内置行业指南灌库）
说明：
    - 集合空间为 cosine，查询距离直接转相似度（1 - distance）。
    - 所有方法尽力而为：向量库不可用时抛 KnowledgeUnavailable，
      由 API 层转业务错误码 1401，不静默降级（聊天需要真实知识）。
    - 用户自定义知识库使用独立集合 user_knowledge，每个片段带
      user_id / doc_id 元数据，检索与删除都强制按 user_id 隔离。
====================================================================
"""

from __future__ import annotations

import io  # docx 内存读取
import logging  # 标准库日志
import re  # 段落/句子切分
import uuid  # 随机文档/片段 ID
from collections import Counter  # 分类统计
from datetime import datetime  # MinIO 路径日期
from pathlib import Path  # 文件名安全处理
from typing import Any  # 泛型类型

import chromadb  # 向量库客户端

from app.config import settings  # Chroma/MinIO 配置
from app.services.embeddings import get_embedder  # 共享 embedding 单例
from app.services.storage import storage  # MinIO 原文备份

logger = logging.getLogger(__name__)  # 模块级日志器

_COLLECTION_NAME = "docagent_knowledge"  # 行业知识库集合
_TEMPLATE_COLLECTION = "doc_templates"  # 模板集合（聊天时补充检索）
_USER_COLLECTION = "user_knowledge"  # 用户自定义知识库集合（多租户，按 user_id 隔离）
_CHUNK_SIZE = 600  # 单个片段目标长度（字符）

_chroma_client: Any | None = None  # Chroma 客户端单例
_collection: Any | None = None  # 知识库集合单例
_template_collection: Any | None = None  # 模板集合单例
_user_collection: Any | None = None  # 用户知识库集合单例


class KnowledgeUnavailable(Exception):
    """知识库不可用（Chroma 连接失败 / 模型加载失败等）。"""


def _get_client() -> Any:
    """获取 Chroma 客户端（lazy 单例，失败抛 KnowledgeUnavailable）。"""
    global _chroma_client
    if _chroma_client is None:
        try:
            _chroma_client = chromadb.HttpClient(
                host=settings.chroma_host, port=settings.chroma_port
            )
        except Exception as exc:  # 连接失败
            raise KnowledgeUnavailable(f"向量库连接失败: {exc}") from exc
    return _chroma_client


def _get_collection() -> Any:
    """获取（并按需创建）知识库集合（cosine 空间）。"""
    global _collection
    if _collection is None:
        try:
            _collection = _get_client().get_or_create_collection(
                name=_COLLECTION_NAME, metadata={"hnsw:space": "cosine"}
            )
        except Exception as exc:  # 创建/获取失败
            raise KnowledgeUnavailable(f"知识库集合不可用: {exc}") from exc
    return _collection


def _get_template_collection() -> Any:
    """获取模板向量集合（聊天时补充检索；不存在返回 None）。"""
    global _template_collection
    if _template_collection is None:
        try:
            _template_collection = _get_client().get_collection(_TEMPLATE_COLLECTION)
        except Exception:  # 模板集合不存在 → None
            _template_collection = None
    return _template_collection


def _get_user_collection() -> Any:
    """获取（并按需创建）用户自定义知识库集合（cosine 空间）。"""
    global _user_collection
    if _user_collection is None:
        try:
            _user_collection = _get_client().get_or_create_collection(
                name=_USER_COLLECTION, metadata={"hnsw:space": "cosine"}
            )
        except Exception as exc:  # 创建/获取失败
            raise KnowledgeUnavailable(f"用户知识库集合不可用: {exc}") from exc
    return _user_collection


def _reset_user_collection() -> None:
    """清空用户知识库集合缓存（集合被删除/重建后重新获取）。"""
    global _user_collection
    _user_collection = None


def _safe_count(collection: Any, label: str) -> int:
    """安全获取集合数量：Chroma 集合被删除/连接异常时统一转 KnowledgeUnavailable。"""
    try:
        return collection.count()
    except Exception as exc:
        raise KnowledgeUnavailable(f"{label}不可用: {exc}") from exc


def chunk_text(text: str, chunk_size: int = _CHUNK_SIZE, overlap: int = 100) -> list[str]:
    """文档切块：按段落聚合，超长段落按句切，块间保留重叠保证上下文连续。

    :param text: 原始文本
    :param chunk_size: 目标片段长度（字符）
    :param overlap: 相邻片段重叠字符数
    :return: 片段列表（去空去重）
    """
    text = re.sub(r"\r\n?", "\n", text or "").strip()
    if not text:
        return []
    paragraphs = [p.strip() for p in text.split("\n") if p.strip()]
    chunks: list[str] = []
    buffer = ""
    for para in paragraphs:
        # 超长段落按句切分（中文句号/感叹号/问号/分号）
        if len(para) > chunk_size:
            sentences = [s.strip() for s in re.split(r"(?<=[。！？!?；;])", para) if s.strip()]
            pieces = sentences
        else:
            pieces = [para]
        for piece in pieces:
            if len(buffer) + len(piece) + 1 <= chunk_size:
                buffer = f"{buffer}\n{piece}".strip()
            else:
                if buffer:
                    chunks.append(buffer)
                buffer = piece
        if len(buffer) >= chunk_size:
            chunks.append(buffer)
            # 尾部保留 overlap 字符作为下一块前缀
            buffer = buffer[-overlap:] if overlap and len(buffer) > overlap else ""
    if buffer.strip():
        chunks.append(buffer.strip())
    # 去重（按前 50 字符指纹）并过滤空片段
    seen: set[str] = set()
    result: list[str] = []
    for c in chunks:
        c = c.strip()
        if not c:
            continue
        key = c[:50]
        if key in seen:
            continue
        seen.add(key)
        result.append(c)
    return result


def extract_text(filename: str | None, data: bytes) -> str:
    """从上传文档抽取纯文本（支持 .docx / .txt / .md / 其他文本）。

    :param filename: 原始文件名（用于判断类型）
    :param data: 文件字节
    :return: 抽取的文本
    """
    name = (filename or "").lower()
    if name.endswith(".docx"):
        try:
            from docx import Document  # python-docx

            doc = Document(io.BytesIO(data))
            parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:  # 表格内容也抽取
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        parts.append(" | ".join(cells))
            return "\n".join(parts)
        except Exception as exc:
            raise KnowledgeUnavailable(f"docx 解析失败: {exc}") from exc
    # 纯文本：依次尝试 utf-8 / gb18030
    for enc in ("utf-8", "gb18030"):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="ignore")


def add_document(
    title: str,
    category: str,
    text: str,
    *,
    template_name: str | None = None,
    source: str = "user_upload",
    doc_ref: str | None = None,
) -> tuple[list[str], int]:
    """切块、向量化并灌入知识库（upsert 幂等）。

    :param title: 文档标题
    :param category: 行业/分类（如"教育"、"商务"）
    :param text: 文档全文
    :param template_name: 关联模板名（可空）
    :param source: 来源（seed / user_upload）
    :param doc_ref: 稳定文档引用（种子灌库时传固定值，保证幂等）
    :return: (片段 ID 列表, 片段数)
    :raises KnowledgeUnavailable: 向量库/模型不可用
    """
    chunks = chunk_text(text)
    if not chunks:
        return [], 0
    try:
        embedder = get_embedder()
        vectors = embedder.encode(chunks, normalize_embeddings=True)
        collection = _get_collection()
        ids: list[str] = []
        documents: list[str] = []
        metadatas: list[dict[str, Any]] = []
        embeddings: list[list[float]] = []
        prefix = "seed" if source == "seed" else "doc"
        for i, chunk in enumerate(chunks):
            if doc_ref:  # 确定性 ID：seed_{ref}_{index}
                cid = f"{prefix}_{doc_ref}_{i:03d}"
            else:
                cid = f"{prefix}_{uuid.uuid4().hex}"
            ids.append(cid)
            documents.append(chunk)
            metadatas.append(
                {
                    "title": title,
                    "category": category,
                    "template_name": template_name or "",
                    "source": source,
                    "chunk_index": i,
                }
            )
            embeddings.append(vectors[i].tolist())
        collection.upsert(ids=ids, documents=documents, metadatas=metadatas, embeddings=embeddings)
        logger.info("[knowledge] 已入库 %d 个片段: %s (%s)", len(chunks), title, category)
        return ids, len(chunks)
    except KnowledgeUnavailable:
        raise
    except Exception as exc:  # 模型/Chroma 异常统一转译
        raise KnowledgeUnavailable(f"向量化失败: {exc}") from exc


def store_original(filename: str, data: bytes) -> str | None:
    """原文备份到 MinIO（失败返回 None，不阻断向量化）。"""
    try:
        safe_name = Path(filename).name  # 防路径穿越
        key = (
            f"knowledge/{datetime.now():%Y/%m/%d}/{uuid.uuid4().hex[:8]}/{safe_name}"
        )
        return storage.upload_bytes(data, bucket=settings.minio_knowledge_bucket, key=key)
    except Exception as exc:  # MinIO 不可用 → 仅告警
        logger.warning("[knowledge] 原文备份失败: %s", exc)
        return None


def add_user_document(
    user_id: int,
    title: str,
    category: str,
    text: str,
) -> tuple[str, int]:
    """用户自定义知识库：切块、向量化并灌入（带 user_id 隔离元数据）。

    :param user_id: 用户主键（元数据 user_id，检索时强制过滤）
    :param title: 文档标题
    :param category: 分类（用户自定义，默认“其他”）
    :param text: 文档全文
    :return: (doc_id, 片段数)
    :raises KnowledgeUnavailable: 向量库/模型不可用
    """
    chunks = chunk_text(text)
    if not chunks:
        return "", 0
    doc_id = uuid.uuid4().hex  # 32 位稳定文档标识
    try:
        embedder = get_embedder()
        vectors = embedder.encode(chunks, normalize_embeddings=True)
        collection = _get_user_collection()
        ids: list[str] = []
        documents: list[str] = []
        metadatas: list[dict[str, Any]] = []
        embeddings: list[list[float]] = []
        for i, chunk in enumerate(chunks):
            cid = f"udoc_{doc_id}_{i:03d}"
            ids.append(cid)
            documents.append(chunk)
            metadatas.append(
                {
                    "user_id": user_id,
                    "doc_id": doc_id,
                    "title": title,
                    "category": category,
                    "source": "user",
                    "chunk_index": i,
                }
            )
            embeddings.append(vectors[i].tolist())
        collection.upsert(
            ids=ids, documents=documents, metadatas=metadatas, embeddings=embeddings
        )
        logger.info(
            "[knowledge] 用户 %s 知识库已入库 %d 个片段: %s (%s)",
            user_id,
            len(chunks),
            title,
            category,
        )
        return doc_id, len(chunks)
    except KnowledgeUnavailable:
        raise
    except Exception as exc:  # 模型/Chroma 异常统一转译
        raise KnowledgeUnavailable(f"用户知识库向量化失败: {exc}") from exc


def search_user(query: str, user_id: int, top_k: int = 5) -> list[dict[str, Any]]:
    """检索用户自己的知识库（强制按 user_id 过滤，互不可见）。

    :param query: 用户问题
    :param user_id: 当前登录用户主键
    :param top_k: 返回片段数
    :return: 片段列表（含 title/category/content/score）
    :raises KnowledgeUnavailable: 向量库/模型不可用
    """
    collection = _get_user_collection()
    try:
        count_result = collection.get(where={"user_id": user_id}, include=[])
    except Exception as exc:
        _reset_user_collection()  # 集合可能已被删除，下次重新获取
        raise KnowledgeUnavailable(f"用户知识库查询失败: {exc}") from exc
    if not count_result.get("ids"):
        return []
    try:
        vec = get_embedder().encode(query, normalize_embeddings=True).tolist()
        result = collection.query(
            query_embeddings=[vec],
            n_results=min(top_k, len(count_result["ids"])),
            where={"user_id": user_id},
            include=["documents", "metadatas", "distances"],
        )
    except KnowledgeUnavailable:
        raise
    except Exception as exc:
        _reset_user_collection()
        raise KnowledgeUnavailable(f"用户知识库检索失败: {exc}") from exc

    hits: list[dict[str, Any]] = []
    for i, cid in enumerate(result["ids"][0]):
        meta = result["metadatas"][0][i] or {}
        hits.append(
            {
                "chunk_id": cid,
                "content": result["documents"][0][i],
                "title": meta.get("title", ""),
                "category": meta.get("category", ""),
                "score": round(1 - result["distances"][0][i], 4),
                "source": "user",
                "doc_id": meta.get("doc_id", ""),
            }
        )
    return hits


def delete_user_document(user_id: int, doc_id: str) -> int:
    """删除用户某篇文档的全部向量片段（按 user_id + doc_id 精准清理）。

    :param user_id: 用户主键
    :param doc_id: 文档标识
    :return: 删除的片段数
    :raises KnowledgeUnavailable: 向量库不可用
    """
    try:
        collection = _get_user_collection()
        count_result = collection.get(
            where={"user_id": user_id, "doc_id": doc_id}, include=[]
        )
        ids = count_result.get("ids") or []
        if ids:
            collection.delete(where={"user_id": user_id, "doc_id": doc_id})
        logger.info(
            "[knowledge] 用户 %s 删除文档 %s 的 %d 个片段",
            user_id,
            doc_id,
            len(ids),
        )
        return len(ids)
    except KnowledgeUnavailable:
        raise
    except Exception as exc:
        _reset_user_collection()
        raise KnowledgeUnavailable(f"用户知识库删除失败: {exc}") from exc


def count_user_chunks(user_id: int) -> int:
    """统计用户知识库的片段总数（供“我的知识库”统计展示）。"""
    collection = _get_user_collection()
    try:
        count_result = collection.get(where={"user_id": user_id}, include=[])
    except Exception as exc:
        _reset_user_collection()
        raise KnowledgeUnavailable(f"用户知识库统计失败: {exc}") from exc
    return len(count_result.get("ids") or [])


def search(query: str, top_k: int = 5) -> list[dict[str, Any]]:
    """语义检索知识库，返回 Top-K 片段（按相似度降序）。

    :param query: 用户问题（如"教育行业应该用哪个模板"）
    :param top_k: 返回片段数
    :return: 片段列表（含 title/category/template_name/content/score）
    :raises KnowledgeUnavailable: 向量库/模型不可用
    """
    collection = _get_collection()
    count = _safe_count(collection, "平台知识库")
    if count == 0:
        return []
    try:
        vec = get_embedder().encode(query, normalize_embeddings=True).tolist()
        result = collection.query(
            query_embeddings=[vec],
            n_results=min(top_k, count),
            include=["documents", "metadatas", "distances"],
        )
    except KnowledgeUnavailable:
        raise
    except Exception as exc:
        raise KnowledgeUnavailable(f"知识检索失败: {exc}") from exc

    hits: list[dict[str, Any]] = []
    for i, cid in enumerate(result["ids"][0]):
        meta = result["metadatas"][0][i] or {}
        hits.append(
            {
                "chunk_id": cid,
                "content": result["documents"][0][i],
                "title": meta.get("title", ""),
                "category": meta.get("category", ""),
                "template_name": meta.get("template_name", ""),
                "score": round(1 - result["distances"][0][i], 4),  # cosine 相似度
                "source": meta.get("source", "knowledge"),
            }
        )
    return hits


def search_templates(query: str, top_k: int = 3) -> list[dict[str, Any]]:
    """补充检索模板集合（doc_templates），聊天回答更完整。"""
    collection = _get_template_collection()
    if collection is None:
        return []
    try:
        count = collection.count()
    except Exception as exc:  # 模板集合不可用 → 返回空（不影响主检索）
        logger.warning("[knowledge] 模板集合计数失败: %s", exc)
        return []
    if count == 0:
        return []
    try:
        vec = get_embedder().encode(query, normalize_embeddings=True).tolist()
        result = collection.query(
            query_embeddings=[vec],
            n_results=min(top_k, count),
            include=["documents", "metadatas", "distances"],
        )
    except Exception as exc:  # 模板集合查询失败 → 返回空（不影响主检索）
        logger.warning("[knowledge] 模板集合检索失败: %s", exc)
        return []
    hits: list[dict[str, Any]] = []
    for i, cid in enumerate(result["ids"][0]):
        meta = result["metadatas"][0][i] or {}
        hits.append(
            {
                "chunk_id": cid,
                "content": result["documents"][0][i],
                "title": meta.get("template_name", ""),
                "category": meta.get("category", ""),
                "template_name": meta.get("template_name", ""),
                "score": round(1 - result["distances"][0][i], 4),
                "source": "template",
            }
        )
    return hits


def stats() -> dict[str, Any]:
    """知识库统计：总片段数 + 按分类分布。"""
    collection = _get_collection()
    total = _safe_count(collection, "平台知识库")
    categories: dict[str, int] = {}
    if total:
        try:
            metas = collection.get(include=["metadatas"])["metadatas"] or []
        except Exception as exc:
            raise KnowledgeUnavailable(f"平台知识库统计失败: {exc}") from exc
        categories = dict(
            Counter(m.get("category", "其他") for m in metas if m)
        )
    return {"total_chunks": total, "categories": categories}
