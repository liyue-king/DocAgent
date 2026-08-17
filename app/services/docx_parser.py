"""
====================================================================
文件用途：docx 文档解析器（DOM 树生成）
====================================================================
作用：
    用 python-docx 解析 .docx 文档，提取段落样式/字体信息，
    生成标准化的 DOM 树结构（供 Planner 决策和 Executor 定位）。
依赖：
    - python-docx（文档读取）
调用方：
    - app/services/docx_editor.py（按段落 ID 执行修改）
    - app/agents/nodes/planner.py（后续：根据 DOM 树规划操作）
输出结构：
    {"paragraphs": [{"id":0, "style":"heading_1", "style_raw":"标题 1",
      "text":"...", "font_name":"宋体", "font_size_pt":12, "bold":False,
      "para_obj":<Paragraph>}], "paragraph_count":10, "tables_count":0}
说明：
    - 样式名做中英文规范化映射（"标题 1"→"heading_1", "Heading 1"→"heading_1" 等）。
    - 保留 python-docx 的原始 Paragraph 对象引用（para_obj），
      编辑器通过段落 id 取值后直接操作原对象，无需二次查找。
====================================================================
"""

from __future__ import annotations

import re  # 正则：样式名映射
from typing import Any  # 泛型类型

from docx import Document  # python-docx 文档对象
from docx.oxml.ns import qn  # 命名空间查询（读取 w:eastAsia 中文字体）
from docx.shared import Length  # 长度对象（行距/段距读取）


def _normalize_style(style_name: str) -> str:
    """将 Word 样式名统一映射为内部键（heading_1|heading_2|heading_3|normal|other）。

    :param style_name: python-docx 返回的 style.name（如 "Heading 1", "标题 1", "Normal"）
    :return: 归一化样式标识
    """
    name = style_name.strip().lower()  # 转小写去空格，方便匹配
    # 匹配中文样式名：标题 1 / 标题 2 / 标题 3 / 正文 / 副标题
    if re.search(r"标题\s*1|heading\s*1|head\s*1", name):
        return "heading_1"
    if re.search(r"标题\s*2|heading\s*2|head\s*2", name):
        return "heading_2"
    if re.search(r"标题\s*3|heading\s*3|head\s*3", name):
        return "heading_3"
    if re.search(r"标题\s*4|heading\s*4|head\s*4", name):
        return "heading_3"  # 四级标题降级为三级
    if re.search(r"正文|normal|body", name):
        return "normal"
    if re.search(r"no spacing|无间距|normal \(web\)|网页正文|list paragraph|列表段落", name):
        return "normal"  # 视觉上等同正文的样式统一按正文处理
    if re.search(r"副标题|subtitle", name):
        return "heading_2"  # 副标题视为二级标题
    if name == "title" or name.endswith("标题"):
        return "heading_1"  # Word "Title"/"标题"（封面大标题常用，此前落到 other 不被处理）
    # 无法匹配的样式归为 other（如列表/引用/题注）
    return "other"


# 常见章节编号（用于把"正文样式+直接格式"的标题识别出来）
_H1_CHAPTER_RE = re.compile(r"^第[0-9一二三四五六七八九十百]+[章节篇卷部分]")
_H1_CN_RE = re.compile(r"^[一二三四五六七八九十]+、")
_H1_DOT_RE = re.compile(r"^[0-9]+[、.．]")
_H2_CN_PAREN_RE = re.compile(r"^（[一二三四五六七八九十0-9]+）")
_H2_DOT_RE = re.compile(r"^[0-9]+\.[0-9]+")
_H3_DOT_RE = re.compile(r"^[0-9]+\.[0-9]+\.[0-9]+")
_DATE_RE = re.compile(r"^[0-9]{4}\s*年")


def _infer_heading_level(
    style_key: str,
    text: str,
    max_size_pt: float,
    any_bold: bool,
) -> str:
    """按内容特征推断标题层级（方案 A：样式名不是标题但内容像标题时识别出来）。

    仅用**编号特征**（第X章 / 一、 / （一） / 1.1 / 1.1.1）做推断——编号是
    结构信息，可信度高。**不再使用外观兜底**（居中/加粗/大字号）：落款、
    强调句与标题视觉特征同构（居中加粗短文本），外观推断必然误判，
    宁可漏判（保持原格式）也不误改。

    :param style_key: 样式名归一化结果
    :param text: 段落纯文本（已 strip）
    :param max_size_pt: 段落内最大字号（磅）
    :param any_bold: 段落内是否存在加粗 run
    :return: heading_1 / heading_2 / heading_3 / 原 style_key
    """
    if style_key in ("heading_1", "heading_2", "heading_3"):
        return style_key  # Word 内置标题样式优先
    if not text or len(text) > 80:
        return style_key  # 空段/长正文不判标题
    if re.search(r"\s{2,}", text):
        return style_key  # 封面字段式排版（多空格占位）不判标题
    if _DATE_RE.match(text):
        return style_key  # 日期（如 2025年11月24日）不判标题
    if len(text) > 60:
        return style_key  # 超过 60 字的段落视为正文，不判标题

    # ---- 编号优先（结构信息，可信）----
    if _H1_CHAPTER_RE.match(text):
        return "heading_1"
    if _H3_DOT_RE.match(text) and (any_bold or max_size_pt >= 12):
        return "heading_3"
    if _H2_DOT_RE.match(text) and (any_bold or max_size_pt >= 12):
        return "heading_2"
    if _H2_CN_PAREN_RE.match(text) and (any_bold or max_size_pt >= 12):
        return "heading_2"
    if _H1_CN_RE.match(text) and (any_bold or max_size_pt >= 14):
        return "heading_1"
    if _H1_DOT_RE.match(text) and (any_bold or max_size_pt >= 14):
        return "heading_1"
    return style_key


def _read_east_asia_font(run: Any) -> str | None:
    """读取 run 的中文字体（w:rFonts/w:eastAsia），未设置返回 None。

    :param run: python-docx Run 对象
    :return: 中文字体名（如"宋体"）；未设置返回 None
    """
    rPr = run._element.rPr  # run 的属性容器（可能为 None）
    if rPr is None:
        return None
    rFonts = rPr.find(qn("w:rFonts"))  # 字体声明元素
    if rFonts is None:
        return None
    return rFonts.get(qn("w:eastAsia"))  # 中文字体名


def _read_line_spacing(para: Any) -> tuple[str | None, float | None]:
    """读取段落行距规则与值。

    :param para: python-docx Paragraph 对象
    :return: (规则字符串, 值)。规则取值 SINGLE/DOUBLE/MULTIPLE/EXACTLY/AT_LEAST；
             SINGLE 值为 1.0，DOUBLE 为 2.0，MULTIPLE 为倍数，EXACTLY/AT_LEAST 为磅值
    """
    pf = para.paragraph_format
    rule_enum = pf.line_spacing_rule  # WD_LINE_SPACING 枚举或 None
    if rule_enum is None:
        return None, None
    rule = rule_enum.name  # 如 "MULTIPLE" / "SINGLE" / "EXACTLY" / "ONE_POINT_FIVE"
    ls = pf.line_spacing  # Length（磅）或 float（倍数）或 None
    # 归一化：Word 内置 1.5/2.5 倍行距在 OOXML 中是 w:lineRule=auto 的
    # 特殊倍数，python-docx 读回为 ONE_POINT_FIVE / TWO_POINT_FIVE，
    # 语义等同模板配置中的 MULTIPLE + 对应倍数值，需归一化后才可比对。
    if rule in ("ONE_POINT_FIVE", "TWO_POINT_FIVE"):
        return "MULTIPLE", (1.5 if rule == "ONE_POINT_FIVE" else 2.5)
    if rule == "SINGLE":
        return rule, 1.0
    if rule == "DOUBLE":
        return rule, 2.0
    if isinstance(ls, Length):  # EXACTLY / AT_LEAST：长度对象（磅）
        return rule, round(ls.pt, 2)
    if ls is not None:  # MULTIPLE：倍数
        return rule, float(ls)
    return rule, None


def _read_paragraph_space(para: Any) -> tuple[float, float]:
    """读取段落段前距与段后距（磅）。

    :param para: python-docx Paragraph 对象
    :return: (space_before_pt, space_after_pt)；未设置视为 0.0
    """
    pf = para.paragraph_format
    sb = pf.space_before
    sa = pf.space_after
    return (round(sb.pt, 2) if sb is not None else 0.0), (
        round(sa.pt, 2) if sa is not None else 0.0
    )


def _read_outline_level(para: Any) -> int | None:
    """读取段落的大纲级别（w:outlineLvl，0-9）。

    Word 最常用的标题实现方式之一："正文样式 + 大纲级别"（样式级或段落级）。
    python-docx 不暴露该属性，必须直读 XML。取值语义：
        0=标题1 … 8=标题9（数字越大级别越低）；9=明确正文文本（body text）。

    查找顺序：段落自身 pPr → 样式链（basedOn 逐级上溯，防循环引用）。

    :param para: python-docx Paragraph 对象
    :return: 大纲级别 int；未设置返回 None
    """

    def _level_of_pPr(pPr: Any) -> int | None:
        if pPr is None:
            return None
        ol = pPr.find(qn("w:outlineLvl"))
        if ol is None:
            return None
        try:
            return int(ol.get(qn("w:val")))
        except (TypeError, ValueError):
            return None  # 非法取值视为未设置

    # 1) 段落自身 pPr（手动设置的大纲级别）
    level = _level_of_pPr(para._element.pPr)
    if level is not None:
        return level
    # 2) 样式链（自定义样式定义的大纲级别）
    seen: set[int] = set()
    style = para.style
    while style is not None and id(style) not in seen:
        seen.add(id(style))
        level = _level_of_pPr(style.element.pPr)
        if level is not None:
            return level
        style = style.base_style  # 沿 basedOn 链上溯（可能为 None）
    return None


def _has_embedded_content(para: Any) -> bool:
    """段落是否含嵌入式内容（图片/形状/公式）。

    这些内容对段落行距敏感：EXACTLY 固定行距下，高于行高的图片/公式会被
    Word 直接裁切显示（不可逆）。检测到后，行距/段距维度应豁免（见
    planner.build_style_ops 与 docx_editor.compute_coverage 的联动）。

    :param para: python-docx Paragraph 对象
    :return: 含 w:drawing（图片/形状）/ w:pict（旧式图片）/ m:oMath（公式）
    """
    tags = (qn("w:drawing"), qn("w:pict"), qn("m:oMath"))
    for el in para._element.iter():
        if el.tag in tags:
            return True
    return False


def build_dom_serial(dom: dict[str, Any]) -> dict[str, Any]:
    """从完整 DOM 提取可序列化部分（剔除 para_obj 引用）。

    供 LangGraph Checkpointer（P1）/ 状态快照持久化使用——para_obj 是
    python-docx 对象引用，无法 JSON 序列化，故单独抽出纯数据层。

    :param dom: build_dom / parse_docx 返回的完整 DOM 树
    :return: 仅含纯数据字段的 DOM（paragraphs 不含 para_obj）
    """
    serial_paragraphs: list[dict[str, Any]] = []
    for p in dom["paragraphs"]:
        serial_paragraphs.append(
            {
                "id": p["id"],
                "style": p["style"],
                "style_raw": p["style_raw"],
                "text": p["text"],
                "font_name": p["font_name"],
                "font_east_asia": p["font_east_asia"],
                "font_size_pt": p["font_size_pt"],
                "bold": p["bold"],
                "any_bold": p["any_bold"],
                "max_font_size_pt": p["max_font_size_pt"],
                "keep_format": p["keep_format"],
                "line_spacing_rule": p["line_spacing_rule"],
                "line_spacing_value": p["line_spacing_value"],
                "space_before_pt": p["space_before_pt"],
                "space_after_pt": p["space_after_pt"],
                "run_count": p["run_count"],
                "outline_lvl": p["outline_lvl"],
                "has_image": p["has_image"],
            }
        )
    return {
        "paragraphs": serial_paragraphs,
        "paragraph_count": dom["paragraph_count"],
        "tables_count": dom["tables_count"],
    }


def parse_docx(file_path: str) -> dict[str, Any]:
    """解析 docx 文件并生成 DOM 树（从文件路径打开）。

    :param file_path: .docx 文件的本地路径
    :return: DOM 树字典
    """
    return build_dom(Document(file_path))


def build_dom(doc: Document) -> dict[str, Any]:
    """基于已有 Document 实例构建 DOM 树（para_obj 与传入 doc 共享引用）。

    与 parse_docx 的区别：不重新打开文件，para_obj 直接引用传入 doc.paragraphs。
    这样 apply_template 修改 DOM 中的 para_obj 时，原 doc 对象也同步更新。

    :param doc: python-docx Document 实例
    :return: DOM 树字典
    """
    paragraphs: list[dict[str, Any]] = []  # 段落列表

    for idx, para in enumerate(doc.paragraphs):
        # 获取样式名（可能为 None）
        style_name = para.style.name if para.style and para.style.name else ""
        style_key = _normalize_style(style_name)

        # 提取段落文字（取纯文本）
        text = para.text.strip()

        # ----- 提取字体信息：取第一个 run 的格式（多数段落只有 1 个 run）-----
        max_font_size_pt = 0.0
        any_bold = False
        bold_run_len = 0  # 加粗 run 的累计文本长度（多 run 段落主格式判定用）
        text_len = 0  # 段落全部 run 的文本总长
        if para.runs:
            run = para.runs[0]  # 首个 run 代表段落主体格式
            font_name = run.font.name  # 西文字体
            font_size = run.font.size
            font_size_pt = round(font_size.pt, 1) if font_size else 0  # 字号（磅）
            font_east_asia = _read_east_asia_font(run)  # 中文字体（w:eastAsia）
            for r in para.runs:
                text_len += len(r.text)
                if r.font.bold:
                    any_bold = True
                    bold_run_len += len(r.text)
                if r.font.size is not None:
                    max_font_size_pt = max(max_font_size_pt, round(r.font.size.pt, 1))
        else:
            font_name = None
            font_size_pt = 0
            bold = False
            font_east_asia = None
        if max_font_size_pt == 0:
            max_font_size_pt = font_size_pt

        # 段落加粗 = 按文本权重取多数（多 run 段落的段首加粗引导语等局部强调
        # 不判整段加粗，避免整段统一覆盖毁掉段内强调格式）
        if para.runs:
            if text_len:
                bold = bold_run_len * 2 >= text_len
            else:
                bold = any_bold
        else:
            bold = False

        # ---- 大纲级别（Word 结构化信息，最高优先）----
        outline_lvl = _read_outline_level(para)
        has_image = _has_embedded_content(para)
        if outline_lvl is not None and outline_lvl <= 8:
            # 0→heading_1、1→heading_2、2→heading_3；3~8（四级及以上）降级三级
            if outline_lvl <= 2:
                style_key = f"heading_{outline_lvl + 1}"
            else:
                style_key = "heading_3"
        elif outline_lvl == 9:
            pass  # 明确正文（body text）：尊重文档声明，不走编号启发式
        else:
            # 智能标题识别：样式名不是标题但内容像标题（仅编号特征，无外观兜底）
            style_key = _infer_heading_level(
                style_key, text, max_font_size_pt, any_bold
            )
        # 保留原格式标记：正文样式的段落原本有加粗（封面字段/强调）→ 模板不强覆盖
        keep_format = style_key == "normal" and any_bold

        # ----- 提取行距与段间距（Validator 五项校验的另两维数据源）-----
        line_spacing_rule, line_spacing_value = _read_line_spacing(para)
        space_before_pt, space_after_pt = _read_paragraph_space(para)

        # 构造段落节点
        node = {
            "id": idx,  # 段落序号（0-based）
            "style": style_key,  # 归一化样式标识
            "style_raw": style_name,  # 原始样式名（调试用）
            "text": text[:200],  # 文本前 200 字（预览用）
            "font_name": font_name,  # 西文字体名（如"Calibri"）
            "font_east_asia": font_east_asia,  # 中文字体名（如"宋体"，可为 None）
            "font_size_pt": font_size_pt,  # 字号（磅）
            "bold": bold,  # 是否加粗
            "any_bold": any_bold,  # 段落内是否有任意加粗 run（保留格式判定）
            "max_font_size_pt": max_font_size_pt,  # 段落内最大字号（标题识别用）
            "keep_format": keep_format,  # 保留原格式（模板不覆盖）
            "line_spacing_rule": line_spacing_rule,  # 行距规则（SINGLE/MULTIPLE/EXACTLY...）
            "line_spacing_value": line_spacing_value,  # 行距值（倍数或磅值）
            "space_before_pt": space_before_pt,  # 段前距（磅）
            "space_after_pt": space_after_pt,  # 段后距（磅）
            "run_count": len(para.runs),  # run 数量（0=空段落，无可排版内容）
            "outline_lvl": outline_lvl,  # 大纲级别（0-9，None=未设置；结构化标题识别）
            "has_image": has_image,  # 含图片/形状/公式（行距裁切风险，三端豁免行距维度）
            "para_obj": para,  # python-docx 段落对象引用（编辑器操作）
        }
        paragraphs.append(node)

    # 统计基本信息
    tables_count = len(doc.tables)  # 表格数量（P0 不处理表格）

    return {
        "paragraphs": paragraphs,
        "paragraph_count": len(paragraphs),
        "tables_count": tables_count,
    }
