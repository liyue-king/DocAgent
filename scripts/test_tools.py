"""
====================================================================
Tools 层综合测试：docx 解析 → 修改 → 覆盖率 + RAG 检索
====================================================================
运行：PYTHONPATH=. python scripts/test_tools.py
====================================================================
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt

from app.models import LogLevel, TaskStatus
from app.services.docx_editor import apply_template, compute_coverage
from app.services.docx_parser import parse_docx


def make_test_docx(file_path: str) -> None:
    """生成测试 docx（含多种样式，模拟真实文档）。"""
    doc = Document()

    # 一级标题：默认无样式，用 built-in heading style
    h1 = doc.add_heading("第一章 项目背景", level=1)
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(12)

    # 正文段落（add_paragraph 副作用即入文档，无需持有引用）
    doc.add_paragraph("这是正文段落，用于测试样式修改功能。")

    # 二级标题
    h2 = doc.add_heading("1.1 技术方案", level=2)
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(6)

    doc.add_paragraph("技术方案包括前端 Vue3 和后端 FastAPI。")

    # 三级标题
    h3 = doc.add_heading("1.1.1 系统架构", level=3)
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(3)

    doc.add_paragraph("系统采用五层架构设计。")

    # 还有一个正文段
    doc.add_paragraph("总结：本文档用于工具层验证测试。")

    doc.save(file_path)


def test_parser(file_path: str) -> dict:
    """测试 docx 解析器。"""
    dom = parse_docx(file_path)
    assert dom["paragraph_count"] >= 4, f"段落数 {dom['paragraph_count']} < 4"
    styles = {p["style"] for p in dom["paragraphs"]}
    assert "heading_1" in styles, f"缺 heading_1: {sorted(styles)}"
    assert "heading_2" in styles, f"缺 heading_2: {sorted(styles)}"
    assert "heading_3" in styles, f"缺 heading_3: {sorted(styles)}"
    assert "normal" in styles, f"缺 normal: {sorted(styles)}"
    print(f"  [PASS] 解析器：{dom['paragraph_count']} 段，样式={sorted(styles)}")
    return dom


def test_editor(doc: Document, config: dict) -> None:
    """测试样式修改器 + 覆盖率。"""
    # 执行修改（apply_template 内部 build_dom，修改即时生效）
    _, ops = apply_template(doc, config)
    print(f"  [INFO] 生成 {len(ops)} 条原子操作")

    # 在同一个 doc 上计算覆盖率
    cov = compute_coverage(doc, config)
    print(f"  [INFO] 覆盖率={cov['coverage']:.2%} passed={cov['passed']}")
    assert cov["coverage"] >= 0.98, f"覆盖率 {cov['coverage']} < 0.98"
    print(f"  [PASS] 编辑器：覆盖率 {cov['coverage']:.2%} ≥ 98%")


def test_models_coherence() -> None:
    """模型内部一致性验证（避免循环导入/枚举值）。"""
    assert TaskStatus.PENDING.value == "pending"
    assert TaskStatus.SUCCESS.value == "success"
    assert LogLevel.INFO.value == "INFO"
    print("  [PASS] 模型枚举一致")


def main() -> None:
    print("Tools 层综合测试")
    print("=" * 50)

    # 种子模板中拿一个来测
    import json

    seed_path = Path(__file__).resolve().parent / "seed_templates.json"
    seeds = json.loads(seed_path.read_text(encoding="utf-8"))
    acad_config = seeds[0]["config"]

    # 生成测试 docx
    file_path = str(
        Path(__file__).resolve().parent.parent / "data" / "_test_tools.docx"
    )
    Path(file_path).parent.mkdir(exist_ok=True)
    make_test_docx(file_path)

    # 解析（返回值仅供 test_parser 内部断言，此处无需持有）
    test_parser(file_path)

    # 编辑器测试（用学术论文模板修改内存中的 doc）
    doc = Document(file_path)
    test_editor(doc, acad_config)

    # 模型一致性
    test_models_coherence()

    print("=" * 50)
    print("Tools 层综合测试全部通过")


if __name__ == "__main__":
    main()
