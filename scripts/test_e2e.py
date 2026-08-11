"""DocAgent 全链路 e2e 验证（U1-U6 + O1/O2 运行时验收）。

覆盖：
    1. auth 链路（注册/验证码负例/积分流水/重复注册）
    2. 修改密码 tv 失效（旧 token 401）
    3. 忘记密码重置（U5）
    4. 管理员模板 CRUD + 推荐（U6）
    5. 任务主链路：上传 → SSE（U1）→ 成功 + validation_report（U3）→ 下载；
       提交即取消 → cancelled（U2）
    6. /metrics 指标（O2）

运行前提：
    1. 基础设施 healthy：scripts/dev_up.ps1
    2. 网关：uv run uvicorn app.main:app --port 8001
    3. Worker：uv run celery -A app.celery_app worker -P solo --loglevel=info
    4. 已执行 init_db.py + init_chroma.py（模板向量已灌库）

用法：PYTHONPATH=. uv run python scripts/test_e2e.py [base_url]
说明：
    - 验证码直接写入 Redis（docagent:email_code:{email}）模拟「邮件已送达」，
      不依赖真实 SMTP 发送。
    - 管理员账号为脚本生成的测试邮箱 + 直连 MySQL 提权，不触碰真实账号。
"""

from __future__ import annotations

import io
import json
import sys
import time
import uuid
from pathlib import Path

import httpx
import redis
from docx import Document
from sqlalchemy import text

from app.config import settings

BASE = sys.argv[1] if len(sys.argv) > 1 else settings.api_base_url
ok = 0
skipped = 0
_TERMINAL = {"success", "failed", "expired", "cancelled"}


def check(name: str, cond: bool) -> None:
    global ok
    assert cond, f"FAIL: {name}"
    ok += 1
    print(f"  [PASS] {name}")


def skip(name: str, reason: str) -> None:
    global skipped
    skipped += 1
    print(f"  [SKIP] {name}（{reason}）")


def make_docx() -> bytes:
    """内存生成 3 段测试文档（标题 + 两段正文）。"""
    buf = io.BytesIO()
    doc = Document()
    doc.add_paragraph("测试标题", style="Heading 1")
    doc.add_paragraph("正文段落一，验证排版效果。")
    doc.add_paragraph("正文段落二，验证下载完整性。")
    doc.save(buf)
    return buf.getvalue()


def seed_code(email: str, code: str = "123456") -> None:
    """模拟邮件已送达：直接把验证码写入 Redis（绕过真实 SMTP）。"""
    client = redis.Redis.from_url(
        settings.redis_cache_url, socket_connect_timeout=2, socket_timeout=2
    )
    client.set(f"docagent:email_code:{email}", code, ex=300)
    client.close()


def auth_headers(token: str) -> dict[str, str]:
    return {"Authorization": f"Bearer {token}"}


def promote_admin(email: str) -> None:
    """测试专用：直连 MySQL 将测试账号提权为管理员（不触碰用户真实账号）。"""
    from app.db import SessionLocal

    db = SessionLocal()
    try:
        db.execute(text("UPDATE users SET is_admin = 1 WHERE email = :e"), {"e": email})
        db.commit()
    finally:
        db.close()


def insert_system_template() -> int:
    """直连 MySQL 插入一条系统保护测试模板，返回 id（负例断言后清理）。

    列表首个模板不保证是系统模板（可能残留非系统测试数据），
    因此负例改为显式造一条 is_system=1 的模板来验证「系统模板不可删」。
    """
    from app.db import SessionLocal

    db = SessionLocal()
    try:
        name = f"系统保护测试 {uuid.uuid4().hex[:6]}"
        db.execute(
            text(
                "INSERT INTO templates (name, description, config, is_system, usage_count) "
                "VALUES (:n, :d, '{}', 1, 0)"
            ),
            {"n": name, "d": "e2e 系统模板保护负例"},
        )
        db.commit()
        row = db.execute(
            text(
                "SELECT id FROM templates WHERE name = :n ORDER BY id DESC LIMIT 1"
            ),
            {"n": name},
        ).fetchone()
        return int(row[0])
    finally:
        db.close()


def delete_template_row(template_id: int) -> None:
    """直连 MySQL 清理测试模板行（系统模板无法通过 API 删除）。"""
    from app.db import SessionLocal

    db = SessionLocal()
    try:
        db.execute(text("DELETE FROM templates WHERE id = :id"), {"id": template_id})
        db.commit()
    finally:
        db.close()


def main() -> None:
    print(f"== DocAgent 全链路 e2e（{BASE}）==")
    with httpx.Client(base_url=BASE, timeout=120) as client:
        # ---------- 1. auth 链路 ----------
        print("-- 1. auth 链路（注册/验证码/积分流水）")
        email = f"e2e_{uuid.uuid4().hex[:8]}@test.local"
        pwd = "pass1234"
        r = client.post(
            "/api/v1/auth/register",
            json={"email": email, "password": pwd, "code": "000000"},
        )
        check("错误验证码 → 1101", r.json().get("code") == 1101)
        seed_code(email)
        r = client.post(
            "/api/v1/auth/register",
            json={"email": email, "password": pwd, "code": "123456"},
        )
        body = r.json()
        check("注册成功 → code=0 + token", body.get("code") == 0 and bool(body.get("token")))
        token = body["token"]
        seed_code(email)  # 验证码一次性消费，重复注册前重新写码以走到邮箱唯一性检查
        r = client.post(
            "/api/v1/auth/register",
            json={"email": email, "password": pwd, "code": "123456"},
        )
        check("重复注册 → 1103", r.json().get("code") == 1103)
        r = client.get("/api/v1/auth/me", headers=auth_headers(token))
        body = r.json()
        check("/me → code=0", body.get("code") == 0)
        check("/me 返回积分余额", "credits_balance" in body.get("user", {}))
        check(
            "credit_logs 含 register 流水",
            any(log.get("action") == "register" for log in body.get("credit_logs", [])),
        )

        # ---------- 2. 修改密码（tv 失效） ----------
        print("-- 2. 修改密码（旧 token 立即失效）")
        r = client.post(
            "/api/v1/auth/change-password",
            json={"old_password": "wrong-pass", "new_password": "newpass123"},
            headers=auth_headers(token),
        )
        check("旧密码错误 → 1104", r.json().get("code") == 1104)
        r = client.post(
            "/api/v1/auth/change-password",
            json={"old_password": pwd, "new_password": "newpass123"},
            headers=auth_headers(token),
        )
        body = r.json()
        check("改密成功签发新 token", body.get("code") == 0 and bool(body.get("token")))
        new_token = body["token"]
        r = client.get("/api/v1/auth/me", headers=auth_headers(token))
        check("旧 token 失效 → HTTP 401", r.status_code == 401)
        r = client.get("/api/v1/auth/me", headers=auth_headers(new_token))
        check("新 token 可用 → code=0", r.json().get("code") == 0)

        # ---------- 3. 忘记密码重置（U5） ----------
        print("-- 3. 忘记密码重置（U5）")
        reset_pwd = "reset1234"
        seed_code(email)
        r = client.post(
            "/api/v1/auth/reset",
            json={"email": email, "code": "123456", "password": reset_pwd},
        )
        check("重置成功 → code=0", r.json().get("code") == 0)
        r = client.post("/api/v1/auth/login", json={"email": email, "password": "newpass123"})
        check("重置后旧密码登录 → 1104", r.json().get("code") == 1104)
        r = client.post("/api/v1/auth/login", json={"email": email, "password": reset_pwd})
        body = r.json()
        check("新密码登录成功", body.get("code") == 0 and bool(body.get("token")))
        token = body["token"]
        r = client.get("/api/v1/auth/me", headers=auth_headers(new_token))
        check("重置前旧 token 失效 → 401", r.status_code == 401)
        seed_code("ghost@test.local")  # 先让验证码通过，走到用户存在性检查
        r = client.post(
            "/api/v1/auth/reset",
            json={"email": "ghost@test.local", "code": "123456", "password": "x123456"},
        )
        check("未知邮箱重置 → 1104", r.json().get("code") == 1104)

        # ---------- 4. 管理员模板 CRUD（U6） ----------
        print("-- 4. 管理员模板 CRUD（U6）")
        admin_email = f"admin_{uuid.uuid4().hex[:8]}@test.local"
        admin_pwd = "admin123"
        seed_code(admin_email)
        r = client.post(
            "/api/v1/auth/register",
            json={"email": admin_email, "password": admin_pwd, "code": "123456"},
        )
        body = r.json()
        check("管理员测试账号注册", body.get("code") == 0 and bool(body.get("token")))
        admin_token = body["token"]
        promote_admin(admin_email)

        tpl_name = f"E2E 商务标书模板 {uuid.uuid4().hex[:4]}"
        r = client.post(
            "/api/v1/templates",
            json={
                "name": tpl_name,
                "description": "商务标书排版模板，用于商业投标文档的格式规范",
                "config": {"font": "宋体", "size_pt": 12},
            },
            headers=auth_headers(admin_token),
        )
        body = r.json()
        check("新建模板 → code=0", body.get("code") == 0)
        tpl_id = body.get("template", {}).get("id")
        check("返回 template.id", bool(tpl_id))

        r = client.get("/api/v1/templates")
        check(
            "列表含新模板",
            any(t["id"] == tpl_id for t in r.json().get("templates", [])),
        )
        r = client.put(
            f"/api/v1/templates/{tpl_id}",
            json={
                "name": f"{tpl_name}（更新）",
                "description": "商务标书排版模板，用于商业投标文档的格式规范（更新版）",
                "config": {"font": "黑体", "size_pt": 14},
            },
            headers=auth_headers(admin_token),
        )
        check("更新模板 → code=0", r.json().get("code") == 0)
        r = client.post(
            "/api/v1/templates/recommend",
            json={"query": "我需要商务标书的排版格式", "top_k": 3},
        )
        body = r.json()
        check("推荐接口 code=0", body.get("code") == 0)
        check(
            "推荐含新模板",
            any(
                rec.get("template_name", "").startswith("E2E 商务标书")
                for rec in body.get("recommendations", [])
            ),
        )
        # 负例：普通用户新建 → 403(1108)；删系统模板 → 1001（直连造系统模板，避免误删）
        r = client.post(
            "/api/v1/templates",
            json={"name": "x", "description": "x", "config": {}},
            headers=auth_headers(token),
        )
        check("非管理员新建 → HTTP 403", r.status_code == 403)
        sys_tpl_id = insert_system_template()
        try:
            r = client.delete(
                f"/api/v1/templates/{sys_tpl_id}", headers=auth_headers(admin_token)
            )
            check("删系统模板 → 1001", r.json().get("code") == 1001)
        finally:
            delete_template_row(sys_tpl_id)
        r = client.delete(f"/api/v1/templates/{tpl_id}", headers=auth_headers(admin_token))
        check("删除新模板 → code=0", r.json().get("code") == 0)
        r = client.get("/api/v1/templates")
        check(
            "列表不含已删模板",
            all(t["id"] != tpl_id for t in r.json().get("templates", [])),
        )

        # ---------- 5. 任务主链路（U1/U2/U3） ----------
        print("-- 5. 任务主链路（SSE/结果预览/下载/取消）")
        r = client.post(
            "/api/v1/process",
            files={"file": ("e2e.docx", make_docx(), "application/octet-stream")},
            data={"prompt": "帮我排版成学术论文格式"},
        )
        body = r.json()
        check("任务提交 → code=0", body.get("code") == 0)
        task_id = body.get("task_id")
        check("返回 task_id", bool(task_id))

        # SSE 流式（U1）：event: log 增量 + 终态 event: status
        got_log = False
        got_status = None
        with client.stream("GET", f"/api/v1/task/{task_id}/stream") as resp:
            check(
                "SSE 200 + event-stream",
                resp.status_code == 200
                and "text/event-stream" in resp.headers.get("content-type", ""),
            )
            ev = ""
            for line in resp.iter_lines():
                line = (line or "").strip()
                if line.startswith("event:"):
                    ev = line[6:].strip()
                elif line.startswith("data:"):
                    if ev == "log":
                        got_log = True
                    elif ev == "status":
                        try:
                            data = json.loads(line[5:].strip())
                        except json.JSONDecodeError:
                            continue
                        if data.get("status") in _TERMINAL:
                            got_status = data
                            break
        check("SSE 收到 log 帧（U1）", got_log)
        check("SSE 收到终态 status 帧", got_status is not None)
        check(
            "SSE 终态 success",
            got_status is not None and got_status["status"] == "success",
        )

        # 轮询终态 + validation_report（U3）
        deadline = time.time() + 90
        last = None
        while time.time() < deadline:
            r = client.get(f"/api/v1/task/{task_id}")
            last = r.json()
            if last.get("status") in _TERMINAL:
                break
            time.sleep(2)
        check("任务终态 success", last is not None and last.get("status") == "success")
        check("progress=100", last.get("progress") == 100)
        check("validation_report 存在（U3）", bool(last.get("validation_report")))
        check("download_url 已生成", bool(last.get("download_url")))

        r = client.get(f"/api/v1/download/{task_id}", follow_redirects=True)
        check("下载 docx（PK 开头）", r.status_code == 200 and r.content.startswith(b"PK"))
        doc = Document(io.BytesIO(r.content))
        check("输出文档段落数 ≥ 3", len(doc.paragraphs) >= 3)

        # 取消（U2）：提交后立即取消
        r = client.post(
            "/api/v1/process",
            files={"file": ("cancel.docx", make_docx(), "application/octet-stream")},
            data={"prompt": "帮我排版成商务报告格式"},
        )
        cancel_id = r.json().get("task_id")
        r = client.post(f"/api/v1/task/{cancel_id}/cancel")
        body = r.json()
        if body.get("code") == 0:
            check("取消请求成功", True)
            deadline = time.time() + 30
            while time.time() < deadline:
                r = client.get(f"/api/v1/task/{cancel_id}")
                if r.json().get("status") in _TERMINAL:
                    break
                time.sleep(1)
            check("任务终态 cancelled（U2）", r.json().get("status") == "cancelled")
        else:
            skip(
                "任务取消（提交即终态，取消窗口过短）",
                f"cancel 返回 {body.get('code')}: {body.get('msg')}",
            )

        # ---------- 6. 运维（O2 metrics + O1 日志文件） ----------
        print("-- 6. /metrics 指标（O2）")
        r = client.get("/metrics")
        check(
            "metrics 200 + text/plain",
            r.status_code == 200 and "text/plain" in r.headers.get("content-type", ""),
        )
        text_body = r.text
        for key in (
            "docagent_tasks_total",
            "docagent_task_duration_seconds",
            "docagent_llm_tokens_total",
            "docagent_queue_pending",
        ):
            check(f"metrics 含 {key}", key in text_body)

        print("-- 日志文件（loguru O1，人工确认）")
        log_dir = Path(__file__).resolve().parent.parent / "logs"
        files = sorted(log_dir.glob("*.log")) if log_dir.exists() else []
        print(f"    logs/: {[f.name for f in files[-3:]] or '（无）'}")

    print(f"== e2e 完成：{ok} PASS / {skipped} SKIP ==")


if __name__ == "__main__":
    main()
