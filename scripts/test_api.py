"""API 网关 e2e 冒烟测试：上传 → 轮询 → 下载全链路。

运行前提：
    1. 基础设施已启动：docker compose up -d
    2. 网关已启动：uv run uvicorn app.main:app --port 8001
    3. Worker 已启动：uv run celery -A app.celery_app worker -P solo --loglevel=info

用法：PYTHONPATH=. uv run python scripts/test_api.py [base_url]
"""

import io
import sys
import time
import uuid

import httpx
from docx import Document

from app.config import settings

# 网关端口默认 8001：8000 被 Chroma 容器占用（docker-compose 映射），不可复用
BASE = sys.argv[1] if len(sys.argv) > 1 else settings.api_base_url
ok = 0


def check(name: str, cond: bool) -> None:
    global ok
    assert cond, f"FAIL: {name}"
    ok += 1
    print(f"  [PASS] {name}")


def make_docx() -> bytes:
    """内存生成 3 段测试文档（标题 + 两段正文）。"""
    buf = io.BytesIO()
    doc = Document()
    doc.add_paragraph("测试标题", style="Heading 1")
    doc.add_paragraph("正文段落一，验证排版效果。")
    doc.add_paragraph("正文段落二，验证下载完整性。")
    doc.save(buf)
    return buf.getvalue()


def main() -> None:
    print(f"== DocAgent API e2e（{BASE}）==")
    with httpx.Client(base_url=BASE, timeout=settings.api_client_timeout_seconds) as client:
        # ---- 1. 健康检查 ----
        r = client.get("/api/v1/health")
        check("health 返回 200", r.status_code == 200)
        body = r.json()
        check(
            "health 服务探测结构",
            set(body["services"]) == {"mysql", "redis", "minio", "chroma"},
        )
        print(f"    services: {body['services']}")

        # ---- 2. 负例：参数校验 ----
        r = client.post(
            "/api/v1/process",
            files={"file": ("x.txt", b"hi", "text/plain")},
            data={"prompt": "排版"},
        )
        check("非 docx → 1003", r.json().get("code") == 1003)
        r = client.post(
            "/api/v1/process",
            files={"file": ("x.docx", make_docx(), "application/octet-stream")},
            data={"prompt": "  "},
        )
        check("prompt 为空 → 1001", r.json().get("code") == 1001)
        r = client.get(f"/api/v1/task/{uuid.uuid4()}")
        check("随机 task_id → 2001", r.json().get("code") == 2001)

        # ---- 3. 主链路：上传 → 轮询 → 下载 ----
        r = client.post(
            "/api/v1/process",
            files={"file": ("sample.docx", make_docx(), "application/octet-stream")},
            data={"prompt": "帮我排版成学术论文格式"},
        )
        body = r.json()
        check("提交成功 → code=0", body.get("code") == 0)
        task_id = body.get("task_id")
        check("返回 task_id", bool(task_id))
        print(f"    task_id: {task_id}")

        # 未完成时下载应被拒绝（2002）
        r = client.get(f"/api/v1/download/{task_id}")
        check("未完成下载 → 2002", r.json().get("code") == 2002)

        # 轮询直至终态（≤90s）
        deadline = time.time() + 90
        last = None
        while time.time() < deadline:
            r = client.get(f"/api/v1/task/{task_id}")
            body = r.json()
            last = body
            if body.get("status") in ("success", "failed", "expired"):
                break
            time.sleep(2)
        check(
            "终态到达 success",
            last.get("status") == "success",
        )
        check("进度 100", last.get("progress") == 100)
        check("日志非空", bool(last.get("logs")))
        print(
            f"    status: {last.get('status')} | progress: {last.get('progress')} | logs: {len(last.get('logs', []))}"
        )
        print("    step:", last.get("step"))
        check("download_url 已生成", bool(last.get("download_url")))

        # ---- 4. 下载并验证 docx 完整性 ----
        r = client.get(f"/api/v1/download/{task_id}", follow_redirects=True)
        check("下载成功", r.status_code == 200 and r.content.startswith(b"PK"))
        doc = Document(io.BytesIO(r.content))
        check("输出文档段落数 ≥ 3", len(doc.paragraphs) >= 3)
        print(f"    输出段落数: {len(doc.paragraphs)}")

    print(f"== 全部通过（{ok} 项）==")


if __name__ == "__main__":
    main()
