"""文档结构识别增强测试：outlineLvl 直读 / 外观兜底移除 / 图片公式段行距豁免。"""

from __future__ import annotations

import base64
import io

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from app.agents.nodes.planner import build_missed_patch_ops, build_style_ops
from app.services.docx_editor import compute_coverage
from app.services.docx_parser import (
    _infer_heading_level,
    _read_outline_level,
    build_dom,
    build_dom_serial,
)

# 1x1 透明 PNG（add_picture 需要真实图片字节）
_PNG_BYTES = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mP8z8BQDwAEhQGAhKmMIQAAAABJRU5ErkJggg=="
)

_TEMPLATE = {
    "paragraph_styles": {
        "heading_1": {"font_name": "黑体", "font_size_pt": 16, "bold": True,
                      "space_before_pt": 24, "space_after_pt": 12,
                      "line_spacing_rule": "EXACTLY", "line_spacing_value": 28},
        "heading_2": {"font_name": "黑体", "font_size_pt": 14, "bold": True,
                      "space_before_pt": 18, "space_after_pt": 6,
                      "line_spacing_rule": "EXACTLY", "line_spacing_value": 28},
        "heading_3": {"font_name": "黑体", "font_size_pt": 12, "bold": True,
                      "space_before_pt": 12, "space_after_pt": 3,
                      "line_spacing_rule": "EXACTLY", "line_spacing_value": 28},
        "normal": {"font_name": "宋体", "font_size_pt": 12, "bold": False,
                   "space_before_pt": 0, "space_after_pt": 0,
                   "line_spacing_rule": "MULTIPLE", "line_spacing_value": 1.5},
    }
}


def _set_para_outline_lvl(para, val: int) -> None:
    """给段落 pPr 直接追加 w:outlineLvl。"""
    pPr = para._element.get_or_add_pPr()
    ol = pPr.makeelement(qn("w:outlineLvl"), {qn("w:val"): str(val)})
    pPr.append(ol)


# ==================== outlineLvl 识别 ====================


def test_outline_lvl_paragraph_level_maps_to_heading() -> None:
    """段落级 outlineLvl=0 → heading_1（正文样式也能识别）。"""
    doc = Document()
    p = doc.add_paragraph("市场分析章节标题")
    _set_para_outline_lvl(p, 0)
    doc.add_paragraph("普通正文。")
    dom = build_dom(doc)
    assert dom["paragraphs"][0]["style"] == "heading_1"
    assert dom["paragraphs"][0]["outline_lvl"] == 0


def test_outline_lvl_style_level_maps_to_heading() -> None:
    """样式级 outlineLvl（自定义样式定义）→ 沿样式链识别。"""
    doc = Document()
    s = doc.styles.add_style("我的章节标题", 1)
    s.base_style = doc.styles["Normal"]
    sPr = s.element.get_or_add_pPr()
    ol = sPr.makeelement(qn("w:outlineLvl"), {qn("w:val"): "1"})
    sPr.append(ol)
    doc.add_paragraph("样式级大纲标题", style=s)
    dom = build_dom(doc)
    assert dom["paragraphs"][0]["style"] == "heading_2"


def test_outline_lvl_9_means_body_text() -> None:
    """outlineLvl=9（明确正文）→ 不被编号启发式提升。"""
    doc = Document()
    p = doc.add_paragraph("一、研究背景")
    _set_para_outline_lvl(p, 9)
    dom = build_dom(doc)
    assert dom["paragraphs"][0]["style"] == "normal"  # 尊重文档声明


def test_outline_lvl_3_plus_downgrades_to_h3() -> None:
    """outlineLvl=3（四级标题）→ 降级 heading_3。"""
    doc = Document()
    p = doc.add_paragraph("四级标题内容")
    _set_para_outline_lvl(p, 3)
    dom = build_dom(doc)
    assert dom["paragraphs"][0]["style"] == "heading_3"


def test_read_outline_level_style_chain_no_cycle() -> None:
    """样式链读取：basedOn 循环引用不崩溃（防呆）。"""
    doc = Document()
    s1 = doc.styles.add_style("循环样式A", 1)
    s2 = doc.styles.add_style("循环样式B", 1)
    s1.base_style = s2
    s2.base_style = s1  # 人为制造循环
    p = doc.add_paragraph("循环样式段落", style=s1)
    assert _read_outline_level(p) is None  # 无 outlineLvl → None，不崩溃


# ==================== 外观兜底移除（落款不误判） ====================


def test_centered_bold_signature_not_heading() -> None:
    """居中加粗 14pt 落款 → 不再误判 heading_1（外观兜底已移除）。"""
    doc = Document()
    p = doc.add_paragraph("某某科技有限公司")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.runs[0]
    r.font.size = Pt(14)
    r.font.bold = True
    dom = build_dom(doc)
    assert dom["paragraphs"][0]["style"] == "normal"


def test_numbered_heading_still_inferred() -> None:
    """编号启发式保留：'第二章 xx' 无条件识别，'一、xx' 加粗 14pt 识别。"""
    assert _infer_heading_level("normal", "第二章 项目背景", 12.0, False) == "heading_1"
    assert _infer_heading_level("normal", "一、研究背景", 14.0, True) == "heading_1"
    assert _infer_heading_level("normal", "（一）数据来源", 12.0, True) == "heading_2"
    assert _infer_heading_level("normal", "1.1 技术方案", 12.0, True) == "heading_2"


# ==================== 图片/公式段检测 ====================


def test_has_image_detected_for_inline_picture() -> None:
    """含 inline 图片的段落 → has_image=True。"""
    doc = Document()
    doc.add_paragraph("普通正文。")
    doc.add_picture(io.BytesIO(_PNG_BYTES), width=Inches(1))
    dom = build_dom(doc)
    assert dom["paragraphs"][1]["has_image"] is True
    assert dom["paragraphs"][0]["has_image"] is False


def test_serial_keeps_new_fields() -> None:
    """build_dom_serial 序列化保留 outline_lvl / has_image。"""
    doc = Document()
    p = doc.add_paragraph("大纲标题")
    _set_para_outline_lvl(p, 0)
    serial = build_dom_serial(build_dom(doc))
    assert serial["paragraphs"][0]["outline_lvl"] == 0
    assert serial["paragraphs"][0]["has_image"] is False


# ==================== 三端联动：行距豁免 ====================


def test_style_ops_skip_spacing_for_image_para() -> None:
    """build_style_ops：含图段落不生成 line_spacing/paragraph_space 指令。"""
    doc = Document()
    doc.add_heading("第一章 图片章", level=1)  # id=0 纯文字
    doc.add_picture(io.BytesIO(_PNG_BYTES), width=Inches(1))  # id=1 含图（heading? no）
    serial = build_dom_serial(build_dom(doc))
    # 图片段落在 heading_1 下会因 _H1_CHAPTER_RE 成 heading_1？——图片段文本为空
    # 直接构造 DOM 保证含图段落样式为 normal 参与分组
    serial["paragraphs"][1]["style"] = "normal"
    ops = build_style_ops(_TEMPLATE, serial)
    spacing_ops = [o for o in ops if o["action"] in ("set_line_spacing", "set_paragraph_space")]
    normal_spacing = [o for o in spacing_ops if 1 in o["para_ids"]]
    assert normal_spacing == []  # 含图段落被剔除出行距/段距指令
    # 纯文字段落（id=0）仍生成行距指令
    assert any(0 in o["para_ids"] for o in spacing_ops)


def test_missed_patch_skips_spacing_for_image_para() -> None:
    """build_missed_patch_ops：含图段落跳过 line_spacing 修补。"""
    report = {
        "missed": [
            {"para_id": 0, "style": "normal", "reason": "font,line_spacing"},
            {"para_id": 1, "style": "normal", "reason": "font,line_spacing"},
        ]
    }
    serial = {
        "paragraphs": [
            {"id": 0, "has_image": False},
            {"id": 1, "has_image": True},
        ]
    }
    ops = build_missed_patch_ops(_TEMPLATE, report, serial)
    dims_for_1 = [o["action"] for o in ops if 1 in o["para_ids"]]
    assert "set_line_spacing" not in dims_for_1
    assert "set_font" in dims_for_1  # 其他维度照常修补
    assert any("set_line_spacing" == o["action"] and 0 in o["para_ids"] for o in ops)


def test_coverage_exempts_spacing_for_image_para() -> None:
    """compute_coverage：纯图片段整段不参与评估（防白跑重试闭环）。"""
    doc = Document()
    doc.add_paragraph("正文文字段落")
    doc.add_picture(io.BytesIO(_PNG_BYTES), width=Inches(1))
    image_para = doc.paragraphs[-1]  # add_picture 返回 InlineShape，段落取最后一个
    dom = build_dom(doc)
    assert dom["paragraphs"][1]["has_image"] is True
    # 给图片段设置行距，使其与模板 EXACTLY/MULTIPLE 不匹配（否则测不出豁免）
    image_para.paragraph_format.line_spacing = Pt(18)
    image_para.paragraph_format.line_spacing_rule = 2  # EXACTLY
    report = compute_coverage(doc, _TEMPLATE)
    image_misses = [m for m in report["missed"] if m["para_id"] == 1]
    assert image_misses == []  # 纯图片段整段豁免 → 不 miss


def test_coverage_exempts_only_spacing_for_caption_para() -> None:
    """带图题文字的图片段：只豁免行距，字体维度照常评估（防重试闭环）。"""
    doc = Document()
    p = doc.add_paragraph("图1-1 系统架构示意图")
    p.add_run().add_picture(io.BytesIO(_PNG_BYTES), width=Inches(1))
    image_para = doc.paragraphs[-1]
    image_para.paragraph_format.line_spacing = Pt(18)
    image_para.paragraph_format.line_spacing_rule = 2  # EXACTLY（与模板 MULTIPLE 1.5 不匹配）
    image_para.runs[0].font.size = Pt(12)  # 图题文字字号对齐模板 normal 12pt
    dom = build_dom(doc)
    assert dom["paragraphs"][0]["has_image"] is True
    assert dom["paragraphs"][0]["text"]  # 有图题文字
    report = compute_coverage(doc, _TEMPLATE)
    miss = next((m for m in report["missed"] if m["para_id"] == 0), None)
    assert miss is None or "line_spacing" not in (miss.get("reason") or "")
