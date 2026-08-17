"""LLM 个性化覆盖测试（v6.2）：用户需求成为验收基准，防重试闭环吞掉覆盖。"""

from __future__ import annotations

from docx import Document

from app.agents.nodes.planner import (
    _llm_augment,
    _validate_llm_op,
    build_missed_patch_ops,
)
from app.services.docx_editor import (
    _apply_overrides,
    apply_operations,
    compute_coverage,
)
from app.services.docx_parser import build_dom, build_dom_serial

_TEMPLATE = {
    "paragraph_styles": {
        "heading_1": {"font_name": "黑体", "font_size_pt": 16, "bold": True,
                      "space_before_pt": 24, "space_after_pt": 12,
                      "line_spacing_rule": "MULTIPLE", "line_spacing_value": 1.5},
        "normal": {"font_name": "宋体", "font_size_pt": 12, "bold": False,
                   "space_before_pt": 0, "space_after_pt": 0,
                   "line_spacing_rule": "MULTIPLE", "line_spacing_value": 1.5},
    }
}


def _doc() -> Document:
    doc = Document()
    doc.add_paragraph("正文段落一。")
    doc.add_paragraph("正文段落二。")
    return doc


# ==================== 核心：LLM 覆盖成为验收基准 ====================


def test_llm_override_is_validation_basis() -> None:
    """核心修复：用户要求楷体 → 验收按楷体（不再被模板宋体判 miss）。"""
    doc = _doc()
    dom = build_dom(doc)
    # 先应用模板全量队列（5 维），再应用 LLM 覆盖指令（楷体）
    from app.agents.nodes.planner import build_style_ops

    apply_operations(dom, build_style_ops(_TEMPLATE, build_dom_serial(dom)))
    apply_operations(dom, [{"action": "set_font", "para_ids": [0, 1], "font": "楷体"}])
    report = compute_coverage(
        doc, _TEMPLATE, {0: {"font": "楷体"}, 1: {"font": "楷体"}}
    )
    assert report["coverage"] >= 1.0  # 修正前：按模板宋体 → 全 miss


def test_llm_override_unmet_still_misses() -> None:
    """用户要求未达成仍会 miss（overrides 只是换基准，不是跳过验收）。"""
    doc = _doc()  # 实际仍是宋体
    report = compute_coverage(doc, _TEMPLATE, {0: {"font": "楷体"}})
    miss = next((m for m in report["missed"] if m["para_id"] == 0), None)
    assert miss is not None and "font" in (miss.get("reason") or "")
    assert miss["expected"]["font_name"] == "楷体"  # expected 展示用户目标


def test_missed_patch_uses_user_value() -> None:
    """修补轮：被覆盖段按用户目标值修补（不再修回模板宋体）。"""
    report = {
        "missed": [
            {"para_id": 0, "style": "normal", "reason": "font,line_spacing"},
        ]
    }
    ops = build_missed_patch_ops(
        _TEMPLATE, report, None,
        {0: {"font": "楷体", "line_spacing": "MULTIPLE",
             "line_spacing_rule": "MULTIPLE", "line_spacing_value": 2.0}}
    )
    font_op = next(o for o in ops if o["action"] == "set_font" and 0 in o["para_ids"])
    assert font_op["font"] == "楷体"  # 用户值
    spacing_op = next(o for o in ops if o["action"] == "set_line_spacing")
    assert spacing_op["value"] == 2.0  # 用户值


def test_uncovered_dims_still_use_template() -> None:
    """未覆盖的段维仍按模板验收（partial override）。"""
    from app.agents.nodes.planner import build_style_ops

    doc = _doc()
    dom = build_dom(doc)
    apply_operations(dom, build_style_ops(_TEMPLATE, build_dom_serial(dom)))
    apply_operations(dom, [{"action": "set_font", "para_ids": [0], "font": "楷体"}])
    # 只覆盖 font，行距未覆盖 → 行距按模板 MULTIPLE 1.5 验收
    report = compute_coverage(doc, _TEMPLATE, {0: {"font": "楷体"}})
    assert report["coverage"] >= 1.0


# ==================== 值校验 ====================


def test_validate_rejects_string_bold() -> None:
    """语义反转防护：bold:"false"（字符串）必须拒绝。"""
    ok, reason = _validate_llm_op(
        {"action": "set_bold", "para_ids": [0], "bold": "false"}, 100, set()
    )
    assert not ok and "布尔" in reason


def test_validate_rejects_bad_size() -> None:
    ok, _reason = _validate_llm_op(
        {"action": "set_font_size", "para_ids": [0], "size_pt": -5}, 100, set()
    )
    assert not ok
    ok, _reason = _validate_llm_op(
        {"action": "set_font_size", "para_ids": [0], "size_pt": 999}, 100, set()
    )
    assert not ok


def test_validate_rejects_bad_rule() -> None:
    ok, _reason = _validate_llm_op(
        {"action": "set_line_spacing", "para_ids": [0], "rule": "1.5倍", "value": 1.5},
        100,
        set(),
    )
    assert not ok


def test_validate_rejects_invisible_para() -> None:
    """para_id 超出 LLM 可见范围（summary 截断）→ 拒绝（防盲猜）。"""
    ok, reason = _validate_llm_op(
        {"action": "set_font", "para_ids": [201], "font": "楷体"}, 200, set()
    )
    assert not ok and "可见范围" in reason


def test_validate_rejects_image_spacing() -> None:
    """图片段行距/段距 → 拒绝（裁切风险，与 has_image 豁免口径一致）。"""
    ok, reason = _validate_llm_op(
        {"action": "set_line_spacing", "para_ids": [5], "rule": "MULTIPLE", "value": 1.5},
        100,
        {5},
    )
    assert not ok and "图片/公式" in reason


def test_validate_accepts_good_op() -> None:
    ok, reason = _validate_llm_op(
        {"action": "set_font", "para_ids": [0, 3], "font": "黑体"}, 100, set()
    )
    assert ok and reason == ""


# ==================== _llm_augment 集成 ====================


def _fake_chat_json(data, tokens=100):
    """构造 chat_json mock。"""

    def _fake(system_prompt, user_prompt, temperature=0.0, max_retries=1):
        return {"data": data, "total_tokens": tokens, "model": "mock"}

    return _fake


def test_llm_augment_produces_overrides_and_unmet(monkeypatch) -> None:
    """合法指令 → overrides；非法指令 → unmet；degraded=False。"""
    from app.services import llm as llm_mod

    monkeypatch.setattr(
        llm_mod, "chat_json",
        _fake_chat_json([
            {"action": "set_font", "para_ids": [0], "font": "楷体"},
            {"action": "set_bold", "para_ids": [1], "bold": "false"},  # 值非法
            {"action": "set_alignment", "para_ids": [2], "align": "center"},  # 白名单外
        ]),
    )
    serial = build_dom_serial(build_dom(_doc()))
    extra, tokens, overrides, unmet, degraded = _llm_augment(
        _TEMPLATE, serial, "改成楷体", []
    )
    assert len(extra) == 1 and extra[0]["font"] == "楷体"
    assert overrides == {0: {"font": "楷体"}}
    assert len(unmet) == 2  # 非法值 + 白名单外
    assert degraded is False
    assert tokens == 100


def test_llm_augment_failure_marks_degraded(monkeypatch) -> None:
    """LLM 调用失败 → degraded=True（前端提示个性化需求未生效）。"""
    from app.services import llm as llm_mod

    def _boom(system_prompt, user_prompt, temperature=0.0, max_retries=1):
        raise RuntimeError("LLM down")

    monkeypatch.setattr(llm_mod, "chat_json", _boom)
    serial = build_dom_serial(build_dom(_doc()))
    extra, _tokens, overrides, _unmet, degraded = _llm_augment(_TEMPLATE, serial, "改格式", [])
    assert extra == [] and overrides == {} and degraded is True


def test_llm_augment_non_list_output_degraded(monkeypatch) -> None:
    """LLM 输出非 JSON 数组 → degraded=True。"""
    from app.services import llm as llm_mod

    monkeypatch.setattr(llm_mod, "chat_json", _fake_chat_json({"not": "a list"}))
    serial = build_dom_serial(build_dom(_doc()))
    _, _, _, _, degraded = _llm_augment(_TEMPLATE, serial, "改格式", [])
    assert degraded is True


# ==================== 辅助函数 ====================


def test_apply_overrides_partial_merge() -> None:
    target = {"font_name": "宋体", "font_size_pt": 12, "bold": False,
              "space_before_pt": 0, "space_after_pt": 0}
    merged = _apply_overrides(target, {"font": "楷体"})
    assert merged["font_name"] == "楷体"  # 覆盖
    assert merged["font_size_pt"] == 12  # 未覆盖保持模板
    assert merged["bold"] is False
