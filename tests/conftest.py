"""pytest 共享夹具：SQLite 内存库 + TestClient（无外部依赖，纯单元测试）。"""

from __future__ import annotations

import io
from pathlib import Path

import pytest
from docx import Document
from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker
from sqlalchemy.pool import StaticPool

from app.db import Base, get_db

_TEST_DOCX_BYTES: bytes | None = None


def _make_docx_bytes() -> bytes:
    """生成含多级标题 + 正文的测试 docx（内存字节，模拟真实文档）。"""
    buf = io.BytesIO()
    doc = Document()
    doc.add_heading("第一章 项目背景", level=1)
    doc.add_paragraph("这是正文段落，用于测试样式修改功能。")
    doc.add_heading("1.1 技术方案", level=2)
    doc.add_paragraph("技术方案包括前端 Vue3 和后端 FastAPI。")
    doc.add_heading("1.1.1 系统架构", level=3)
    doc.add_paragraph("系统采用五层架构设计。")
    doc.save(buf)
    return buf.getvalue()


def get_test_docx_bytes() -> bytes:
    """模块级缓存生成一次，供各测试复用。"""
    global _TEST_DOCX_BYTES
    if _TEST_DOCX_BYTES is None:
        _TEST_DOCX_BYTES = _make_docx_bytes()
    return _TEST_DOCX_BYTES


@pytest.fixture()
def docx_bytes() -> bytes:
    """测试 docx 的字节内容（供 API 测试 multipart 上传）。"""
    return get_test_docx_bytes()


@pytest.fixture()
def test_docx_path(tmp_path: Path) -> str:
    """测试 docx 落盘路径（供 docx 工具链测试）。"""
    path = tmp_path / "sample.docx"
    path.write_bytes(get_test_docx_bytes())
    return str(path)


def _patch_sqlite_defaults() -> None:
    """MySQL 专用默认值（CURRENT_TIMESTAMP(3)）替换为 SQLite 兼容形式。

    模型用 server_default=text("CURRENT_TIMESTAMP(3)")，SQLAlchemy 对 raw
    text 不做方言转换，SQLite 建表会报错；此处仅测试环境替换。
    """
    from sqlalchemy import DefaultClause
    from sqlalchemy import text as sa_text

    for table in Base.metadata.tables.values():
        for col in table.columns:
            sd = col.server_default
            if sd is not None and sd.arg is not None and "CURRENT_TIMESTAMP(3)" in str(sd.arg):
                col.server_default = DefaultClause(sa_text("CURRENT_TIMESTAMP"))


from sqlalchemy import BigInteger
from sqlalchemy.dialects.mysql import TINYINT
from sqlalchemy.ext.compiler import compiles


@compiles(TINYINT, "sqlite")
def _compile_tinyint_sqlite(element, compiler, **kw):
    """MySQL TINYINT → SQLite INTEGER（仅测试环境方言渲染）。"""
    return "INTEGER"


@compiles(BigInteger, "sqlite")
def _compile_bigint_sqlite(element, compiler, **kw):
    """BIGINT → SQLite INTEGER：主键自增在 SQLite 仅支持 INTEGER PRIMARY KEY。"""
    return "INTEGER"


@pytest.fixture()
def db_session():
    """SQLite 内存库会话：全部模型建表，测试结束后销毁。"""
    _patch_sqlite_defaults()
    engine = create_engine(
        "sqlite://",
        connect_args={"check_same_thread": False},
        poolclass=StaticPool,
    )
    Base.metadata.create_all(engine)
    SessionLocal = sessionmaker(bind=engine, autoflush=False, expire_on_commit=False)
    session = SessionLocal()
    try:
        yield session
    finally:
        session.close()
        engine.dispose()


@pytest.fixture()
def client(db_session):
    """FastAPI TestClient：get_db 依赖覆盖为 SQLite 会话。"""
    from fastapi.testclient import TestClient

    from app.main import app

    def override_get_db():
        yield db_session

    app.dependency_overrides[get_db] = override_get_db
    try:
        with TestClient(app) as c:
            yield c
    finally:
        app.dependency_overrides.clear()
