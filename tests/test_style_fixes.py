"""docx 解析/编辑器修复回归：Title 样式、多 run 强调保护、CJK 字体槽。"""

from __future__ import annotations

from docx import Document
from docx.oxml.ns import qn

from app.services.docx_editor import apply_template, compute_coverage
from app.services.docx_parser import _normalize_style, build_dom

_CUSTOM_CONFIG = {
    "paragraph_styles": {
        "heading_1": {
            "font_name": "宋体",
            "font_size_pt": 16,
            "bold": False,
            "space_before_pt": 24,
            "space_after_pt": 12,
            "line_spacing_rule": "MULTIPLE",
            "line_spacing_value": 1.5,
        },
        "heading_2": {
            "font_name": "宋体",
            "font_size_pt": 14,
            "bold": True,
            "space_before_pt": 18,
            "space_after_pt": 6,
            "line_spacing_rule": "MULTIPLE",
            "line_spacing_value": 1.5,
        },
        "heading_3": {
            "font_name": "宋体",
            "font_size_pt": 12,
            "bold": True,
            "space_before_pt": 12,
            "space_after_pt": 3,
            "line_spacing_rule": "MULTIPLE",
            "line_spacing_value": 1.5,
        },
        "normal": {
            "font_name": "宋体",
            "font_size_pt": 12,
            "bold": False,
            "space_before_pt": 0,
            "space_after_pt": 0,
            "line_spacing_rule": "MULTIPLE",
            "line_spacing_value": 1.5,
        },
    }
}


def test_normalize_style_title_heading() -> None:
    """Word "Title"/"标题"（封面大标题）→ heading_1（此前落到 other）。"""
    assert _normalize_style("Title") == "heading_1"
    assert _normalize_style("标题") == "heading_1"
    assert _normalize_style("主标题") == "heading_1"
    # 既有映射不受影响
    assert _normalize_style("标题 1") == "heading_1"
    assert _normalize_style("Heading 2") == "heading_2"
    assert _normalize_style("副标题") == "heading_2"
    assert _normalize_style("Normal") == "normal"


def test_mixed_run_dominant_bold_judgement() -> None:
    """多 run 段落：加粗按文本权重取多数（段首加粗引导语不判整段加粗）。"""
    doc = Document()
    p = doc.add_paragraph()
    r1 = p.add_run("注意：")
    r1.bold = True
    r2 = p.add_run("请按照模板要求认真排版，正文内容较长")
    assert r2.bold is not True

    dom = build_dom(doc)
    node = dom["paragraphs"][0]
    assert node["run_count"] == 2
    assert node["any_bold"] is True  # 存在强调 run（keep_format 判定仍用）
    assert node["bold"] is False  # 权重多数非加粗 → 不判整段加粗


def test_mixed_run_emphasis_preserved_when_dominant_matches() -> None:
    """幂等跳过：段落主格式已达标 → 不再整段覆盖（段首加粗引导语保留）。"""
    doc = Document()
    p = doc.add_paragraph(style="Heading 1")
    r1 = p.add_run("第一章 ")
    r1.bold = True
    p.add_run("总体概述与项目背景说明")  # 权重多数非加粗
    # 模板 heading_1 bold=False；段落主格式 bold=False 已达标 → 跳过 set_bold
    apply_template(doc, _CUSTOM_CONFIG)
    assert doc.paragraphs[0].runs[0].bold is True  # 段首强调未被覆盖


def test_mixed_font_emphasis_preserved_when_first_run_matches() -> None:
    """幂等跳过（字体）：首 run 已达标 → 段内其它 run 的字体强调保留。"""
    doc = Document()
    p = doc.add_paragraph()
    r1 = p.add_run("正文内容")
    r1.font.name = "宋体"  # 与模板目标一致
    r2 = p.add_run("【重点强调】")
    r2.font.name = "楷体"  # 段内强调字体

    apply_template(doc, _CUSTOM_CONFIG)
    assert doc.paragraphs[0].runs[1].font.name == "楷体"  # 强调保留
    cov = compute_coverage(doc, _CUSTOM_CONFIG)
    assert cov["coverage"] >= 0.98  # 覆盖判定与执行口径一致，无假未达标


def test_cjk_font_written_to_east_asia_only() -> None:
    """中文字体只写 w:eastAsia，不污染 w:ascii 西文字体槽。"""
    from app.services.docx_editor import _set_paragraph_font

    doc = Document()
    p = doc.add_paragraph()
    r = p.add_run("abc 中文 mixed text")
    _set_paragraph_font(p, "宋体")
    rFonts = r._element.rPr.find(qn("w:rFonts"))
    assert rFonts.get(qn("w:eastAsia")) == "宋体"
    assert rFonts.get(qn("w:ascii")) is None  # 西文字体槽未被中文名污染
    assert r.font.name is None


def test_latin_font_written_to_ascii_normally() -> None:
    """西文字体仍走常规通道（w:ascii + w:hAnsi）。"""
    from app.services.docx_editor import _set_paragraph_font

    doc = Document()
    p = doc.add_paragraph()
    r = p.add_run("hello world")
    _set_paragraph_font(p, "Times New Roman")
    assert r.font.name == "Times New Roman"
